import io
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import asyncio

async def generate_docx(answer, time):
    doc = Document()

    async def split_tags(reply):
        pattern = r'\[(.*?)\](.*?)\[/\1\]'
        tags = re.findall(pattern, reply, re.DOTALL)
        return tags

    async def parse_response(tags_array):
        if not tags_array:
            raise IndexError
        
        # Add title and subtitle
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run("O'zbekiston Respublikasi Raqamli texnologiyalar\nvazirligi.")
        run1.bold = True
        run1.font.size = Pt(24)

        # p2 = doc.add_paragraph()
        # p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run2 = p2.add_run("Muhammad Al-Xorazmiy nomidagi Toshkent\nAxborot Texnologiyalari Universiteti")
        # run2.bold = True
        # run2.font.size = Pt(22)
        university_name = "Muhammad Al-Xorazmiy nomidagi Toshkent\nAxborot Texnologiyalari Universiteti"
        for item in tags_array:
            if item[0] == 'UNIVERSITY':
                university_name = item[1]
                break

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(university_name)
        run2.bold = True
        run2.font.size = Pt(22)

        # Add empty lines
        for _ in range(4):
            doc.add_paragraph()

        # Add "O'rnatilgan tizimlar fanidan"
        theme_text = "O'rnatilgan tizimlar fanidan"
        for item in tags_array:
            if item[0] == 'THEME':
                theme_text = item[1]
                break

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(theme_text)
        run1.font.size = Pt(16)

        doc.add_paragraph()

        # Add "Mustaqil ish"
        p = doc.add_paragraph("Mustaqil ish")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(20)

        # Gather subtitles for the table of contents
        subtitles = []
        right_content = ""
        for item in tags_array:
            if item[0] == 'SUBTITLE':
                subtitles.append(item[1])
            if item[0] == 'RIGHTCONTENT':
                right_content = item[1]  # Store the right content

        # Add "Toshkent 2024"
        if right_content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(right_content)
            run.font.size = Pt(14)

        for _ in range(4):
          doc.add_paragraph()
        
        p_tosh = doc.add_paragraph()
        p_tosh.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tosh1 = p_tosh.add_run("Toshkent 2024")
        p_tosh1.bold = True
        p_tosh1.font.size = Pt(14)

        doc.add_page_break()

        # Add "Reja" (Table of Contents)
        mundarija = doc.add_paragraph("Reja")
        mundarija.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mundarija.runs[0].bold = True
        mundarija.runs[0].font.size = Pt(14)

        doc.add_paragraph()

        # Generate the Table of Contents based on SUBTITLE
        data = []
        page_number = 3  # Start page numbering from 3 (or adjust as needed)
        for i, subtitle in enumerate(subtitles):
            data.append((f"{i + 1}.", subtitle, str(page_number)))
            page_number += 1  # Increment page number for each subtitle

        # Add Table of Contents
        for num, content, page in data:
            paragraph = doc.add_paragraph()
            paragraph.add_run(num + " ")
            paragraph.add_run(content)
            
            paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_ALIGNMENT.DECIMAL)
            paragraph.add_run().add_tab()
            
            paragraph.add_run(page)

        doc.add_page_break()

        # Add content based on the tags_array
        for item in tags_array:
            match (item[0]):
                case 'HEADING':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(item[1])
                    run.bold = True
                    run.font.size = Pt(14)
                case 'CONTENT':
                    p = doc.add_paragraph(item[1])
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # case 'RIGHTCONTENT':
                #     # Add the RIGHTCONTENT section
                #     p = doc.add_paragraph()
                #     p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                #     run = p.add_run(item[1])
                #     run.font.size = Pt(12)
                case 'IMAGE':
                    # Image handling is commented out as it requires additional setup
                    pass

    async def find_title(tags_array):
        for item in tags_array:
            if item[0] == 'TITLE':
                return item[1]
        return "Untitled Document"

    reply_array = await split_tags(answer)
    await parse_response(reply_array)
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_title = f"{await find_title(reply_array)}.docx"
    with open(f"output2{time}.docx", "wb") as f:
        f.write(docx_bytes)
    print(f"Document saved as output2{time}.docx")
    print(f"Done: {docx_title}")

    return docx_bytes, docx_title

# Example usage
# async def main():
#     sample_answer = """[TITLE]Atom tuzilishining asoslari: elektronlar, protonlar va neytronlar.[/TITLE]
# [SUBTITLE]Atom tuzilishining asoslari: elektronlar, protonlar va neytronlar.[/SUBTITLE]
# [HEADING]Atom tuzilishining asoslari: elektronlar, protonlar va neytronlar.[/HEADING]
# [CONTENT]Elektronlar atomning negativ zarur qismini tashkil etadi va atom yadro tashqarisida joylashganlar. Uning massasi yadro partikullari bilan solishtirilganda juda kichik.[/CONTENT]
# [CONTENT]Protonlar atomning positiv zarur qismini tashkil etadi va yadroda joylashganlar. Ular elektronlardan katta massaga ega bo'lib, atomning kimiyoviy xususiyatlarini aniqlashda ahamiyatga ega.[/CONTENT]
# [CONTENT]Neytronlar atomning neutral zarur qismini tashkil etadi va yadroda protonlar bilan birga paydo bo'ladi. Ular atomlarning stabi­li­ya­ti­ni ta'minlashda muhimiy ahamiyatga ega.[/CONTENT]
# [CONTENT]Atom elektron, proton va neytronlardan iborat. Protonlar va neytronlar atomning yadrosida joylashgan bo'lib, elektronlar esa yadro atrofiga aylanadi. Ular atom tuzilishining asoslari hisoblanadi.[/CONTENT]
# [SUBTITLE]Elektron qobig‘ining qatlamlari va ularning atom xususiyatlari uchun ahamiyati.[/SUBTITLE]
# [HEADING]Elektron qobig‘ining qatlamlari va ularning atom xususiyatlari uchun ahamiyati.[/HEADING]
# [CONTENT]Elektron qobig‘ini tashkil etuvchi qatlamlar undagi elektronlar soni va joylashuvi bo‘yicha taqsimlangan. Ular K, L, M, N qatlamli bo‘lib, har bir qatlamin joylashuvi o‘ziga xosdir.[/CONTENT]
# [CONTENT]Elektron qobig‘ining qatlamlari atomning kimyoviy xususiyatlari va uni boshqa elementlar bilan taqqoslashda ahamiyatga ega. Elektron qatlamlarining joylashuvi va umumiy soni atomlar orasidagi kimyoviy aloqalarni ta’minlaydi.[/CONTENT]
# [CONTENT]Elektron qobig‘ining qatlamlari atom tuzilishidagi tarkibiy xususiyatlarning o‘zi mavjud bo‘lgan har bir atomda elektronlarning joylashuvi va sonini belgilaydi. Bu esa atomning kimyoviy xususiyatlarini tushuntiradi.[/CONTENT]
# [CONTENT]Elektron qobig‘ining qatlamlari energiya darajasini belgilaydi va bu, kimyoviy reaksiyalarni, elementning xususiy kimyoviy hususiyatlarini va atomning o‘zining taxminiy kimyoviy reaksiyalarini ta’minlaydi.[/CONTENT]
# [SUBTITLE]Protonlar va neytronlar: Yadroning tuzilishi.[/SUBTITLE]
# [HEADING]Protonlar va neytronlar: Yadroning tuzilishi.[/HEADING]
# [CONTENT]Atom yadrosining qurishi uchun protonlar va neytronlar, ya'ni nukleonlar jamiyati erkin turkumlar sifatida mustahkam bir hatni egallashadi.[/CONTENT]
# [CONTENT]Atom yadrosi protonlar to'plami, neytronlar to'plami va bir-biriga bo'g'li elektronlar to'plamidan iborat. Bu tuzilma nukleonlar va elektronlar orasidagi quvurlar bilan aniqlanadi.[/CONTENT]
# [CONTENT]+1 elektrikli yunon ko'chiruvchi salsa proton deb ataladi. Proton O'tkir tomonga bo'lgan elektronlar orasida sharch tuzilganligi hisoblanadi.[/CONTENT]
# [CONTENT]Neytronlar no'l elektr darajadagi yo'nalishlik fermionlar bilan shunchaki barobar massali, ammo tarqalgan elektr tomon yunon yoki qatlama ko'chiruvlari yo'q.[/CONTENT]
# [SUBTITLE]Atom raqami va massa soni: Atomlarning klassifikatsiyasi.[/SUBTITLE]
# [HEADING]Atom raqami va massa soni: Atomlarning klassifikatsiyasi.[/HEADING]
# [CONTENT]Atom raqami, bu atomning neytron va protonlar sonini ifoda etadi. Massa soni esa atomda mavjud bo'lgan proton va neytronlar yig'indisini anglatadi.[/CONTENT]
# [CONTENT]Atomlar kimyo tarkibiy aniq atomlarga bo'linadi. Ularning har biri o'zlarini kimyolement bo'yicha taxlanishi uchun atom raqam va massasi bilan belgilanadi.[/CONTENT]
# [CONTENT]Atomlar metal, metalloiddan va no-metallardan iborat. Ularning xususiyatlari, kimyoviy reaktsiyalari va o'ziga xos xossalari ularni bir-biriga ajratadi.[/CONTENT]
# [CONTENT]Atomlar proton, neytron va elektrondan iborat. Bu tarkib ularda kimyoviy va fizikaviy xossalarni ta'kidlaydi va ularning davraniyb to'ri.[/CONTENT]
# [SUBTITLE]Kimyoviy elementlar davriy jadvali va atom tuzilishi.[/SUBTITLE]
# [HEADING]Kimyoviy elementlar davriy jadvali va atom tuzilishi.[/HEADING]
# [CONTENT]Kimyoviy elementlar davriy jadvali, elementlar tuzilishiga asoslangan jadvaldir va kimyoviy elementlarni tuzilishi va xususiyatlarini ifodalaydi.[/CONTENT]
# [CONTENT]Atom tuzilishi, kimyoviy elementning yadro, elektron oʻrtasida aloqasini tushuntiradi. Atom tuzilishning asosiy qismlarini tahrir qiladi.[/CONTENT]
# [CONTENT]Kimyoviy elementlar davriy jadvali asosida, atomlar qanday tuzilganligi va bir-biriga qanday aloqada ekanligi ko‘rsatiladi.[/CONTENT]
# [CONTENT]Kimyoviy elementlar xususiyatlari, ularning atomlarining tuziligiga asoslangan, kimyoviy reaksiyalarda qanday namoyon bo‘lishi haqida ma'lumotlar beradi.[/CONTENT]
# [SUBTITLE]Ionlash: Ionlar va ularning atom tuzilishidagi o‘rni.[/SUBTITLE]
# [HEADING]Ionlash: Ionlar va ularning atom tuzilishidagi o‘rni.[/HEADING]
# [CONTENT]Ionlar atom tuzilishidagi o‘rni haqida ma'lumot.[/CONTENT]
# [CONTENT]Ionlar atomdagi elektronlarini yo‘q qilish vositasidir.[/CONTENT]
# [CONTENT]Iyonlar va ularning atom tuzilishidagi o‘rni.[/CONTENT]
# [CONTENT]Ionlar vositasidagi atom elektronlaridan taqaladi.[/CONTENT]
# [SUBTITLE]Kovalent bog‘lanish va atom orbitalarining roli.[/SUBTITLE]
# [HEADING]Kovalent bog‘lanish va atom orbitalarining roli.[/HEADING]
# [CONTENT]Kovalent bog‘lanish, ikki yoki undan ko'p atom o'rtasida elektronlar almashishiga asoslangan molekuli o'rnatish usuli.[/CONTENT]
# [CONTENT]Atom orbital, atomning elektronlar uchun moliyalangan joylarini bildiruvchi matematik qavrayiladi.[/CONTENT]
# [CONTENT]Kovalent bog‘lanishda, atom orbitalari elektronlar almashishini boshlash uchun paydo bo'lishadi va molekul shakl yoki xususiyatlarini belgilashda muhim rol o'ynaydi.[/CONTENT]
# [CONTENT]Kovalent bog‘lanishda, atom orbitalari elektron almashishini muhofaza etadi va molekulning kimyoviy xususiyatlarini belgilashda ahamiyatga ega bo'ladi.[/CONTENT]
# [SUBTITLE]Metallik bog‘lanish va atom tuzilishining metall xususiyatlariga ta’siri.[/SUBTITLE]
# [HEADING]Metallik bog‘lanish va atom tuzilishining metall xususiyatlariga ta’siri.[/HEADING]
# [CONTENT]Metallik bog‘lanish metall ionlarining elektronlarni umumiy ob’yektida (elektron tenglama) qo‘llashiga asoslangan. Uning natijasida metallarning yaxshi issi va elektr uzatish xususiyatlari paydo bo‘ladi.[/CONTENT]
# [CONTENT]Metall ikki yoki undan ko‘p o‘rnatilgan bloklar (qurish, zargarlik, metall, qutb) sifatida to‘plangan atomlar to‘plamidir. Bu tuzilish metallarni elastik, metallarning issi va elektr uzatish qobiliyati bilan ajratadi.[/CONTENT]
# [CONTENT]Metallar bog‘lanish qonuni tufayli yorug‘likka aylanadi. Ular metall sulfat, metall oksit va metall ko‘rinishidagi boshqa kimyoviy modellarga o‘ziga xos bo‘lgan xususiyatlar bilan ajratiladi.[/CONTENT]
# [CONTENT]Metallar kristallarni kristal tuzilishi tufayli tuzishadi. Bu tuzilish metalldagi ganjlik va qattiqlik, termo qatlamlilik, metallarning deformatsiyaga o‘tish qobiliyati va shakl o‘zgarishi to‘g‘risidagi xususiyatlarini ta’minlaydi.[/CONTENT]
# [SUBTITLE]Elektronegativlik va atom tuzilishidagi elektron taqsimlanishi.[/SUBTITLE]
# [HEADING]Elektronegativlik va atom tuzilishidagi elektron taqsimlanishi.[/HEADING]
# [CONTENT]Elektronegativlik - bu atomni elektronlarni jalb qilishga qaratilgan yuqori qobiliyatdir.[/CONTENT]
# [CONTENT]Atomda elektronlar energiya darajalariga joylashgan oʻrnalar bilan taqsimlanadi va ulardan har biri oʻziga xos orbitalda joylashadi.[/CONTENT]
# [CONTENT]Elektronegativlik, kimyoviy elementlar orasidagi atomlar orasidagi elektronlar uchrashishining kuchini ko'rsatadi va Pauling shkalamizda ifodalangan.[/CONTENT]
# [CONTENT]Elektronegativlik, bir elementdan boshqa bir elementga olib kelgan elektronnolloqning sifatlarining farqini aks ettiradi.[/CONTENT]
# [SUBTITLE]Atomning kvant mexanik modeli: Shrodinger tenglamasi.[/SUBTITLE]
# [HEADING]Atomning kvant mexanik modeli: Shrodinger tenglamasi.[/HEADING]
# [CONTENT]Shrodinger tenglamasi, atomning kvant mexanik modeliga asos bo'lgan tenglama, kvant mexanikadagi zaruriy tenglamalardan biridir.[/CONTENT]
# [CONTENT]Atomning kvant mexanik modeli, elektronlar haqida tezlikni bilishni taʼminlovchi kvant mexanika asosida atomning tahlili va energiyasini aniqlash modelidir.[/CONTENT]
# [CONTENT]Shrodinger tenglamasi qonuniy kerakli miqyosli tenglama, atom tuzilmasini va davrining energiya darajasini aniqlashda yordam beradi.[/CONTENT]
# [CONTENT]Kvant mexanik modeli, atomning tahlili va energiyasini aniqlovchi eng to'g'ri va realistlik usuldir, Shrodinger tenglamasi esa bu modelni hisoblashda ishlatiladi.[/CONTENT]
# [SUBTITLE]Atomdagi elektron holatlari va kvant sonlari.[/SUBTITLE]
# [HEADING]Atomdagi elektron holatlari va kvant sonlari.[/HEADING]
# [CONTENT]Atomdagi elektron energiyasiga bog'liq o'zgaruvchan holatlar. Ularning miqdori orbital turi va elektron raqamiga bog'liq bo'ladi.[/CONTENT]
# [CONTENT]Elektronlar energiya, orbital moment, mgnit moment va magnitli kvant sonlar orqali tavsiflanadi. Bu kvant sonlar elektronlarning joylashuvi va davr harakatlarini aniqlaydi.[/CONTENT]
# [CONTENT]Elektronlar atomda fut azotiradagan energiya holatlarida joylashadi. Bu holatlar atomning energetik strukturasi to'g'risida ma'lumot beradi.[/CONTENT]
# [CONTENT]S, P, D va F orbital turlari elektronlar uchun joylashuv joylari sifatida ishlatiladi. Bu orbitalar atomning energetik strukturasi va kimyoviy xossalarni aniqlashda muhimdir.[/CONTENT]
# [SUBTITLE]Valentlik elektronlari va kimyoviy bog‘lanishning kelib chiqishi.[/SUBTITLE]
# [HEADING]Valentlik elektronlari va kimyoviy bog‘lanishning kelib chiqishi.[/HEADING]
# [CONTENT]Valentlik elektronlar - atombardagi eng sirtqi elektronlar. Ular kimyoviy bog‘lanishda o‘z vazifalarini bajaradilar.[/CONTENT]
# [CONTENT]Kimyoviy bog‘lanish asosida Valentlik elektronlar o'zaro bo‘lishishadi. Ushbu elektronlar atom energiyasiga bog‘liqlik ko‘rsatadi.[/CONTENT]
# [CONTENT]Valentlik elektronlar kimyoviy reaktsiyalarda atomlar o‘rtasidagi bog‘lanishni tushunishda muhim ahamiyatga egadir.[/CONTENT]
# [CONTENT]Molekulalar o‘rtasidagi bog‘lanishda Valentlik elektronlar o‘z vazifalarini bajaradi, molekulaning xususiyatlari va kimyoviy xossalari uchun muhimdir.[/CONTENT]
# [SUBTITLE]Atomlarning birikishi: Molekulalar va kristallar.[/SUBTITLE]
# [HEADING]Atomlarning birikishi: Molekulalar va kristallar.[/HEADING]
# [CONTENT]Atomlar kimyoviy moddalar biriktirilganda molekulalar hosil bo'ladi. Ular atomlar orasida kimyoviy aloqalar bilan bog'langanlar. Molekulalar esa o'zaro chimchil qoladi.[/CONTENT]
# [CONTENT]Molekulalar esa kristallar, sular, gazlar yoki boshqa moddalar shaklida bo'lishi mumkin. Ular xususiyatlari molekula tuzilishi va oraliq kuchiga bog'liq.[/CONTENT]
# [CONTENT]Kristallar esa tuzilishi belgilangan va o'zining geometriyasi va tarzda tartiblashiga ega bo'lgan moddalardir. Atomlar kristallarda tartiblangan, to'g'ri ornatingan tarzda joylashadi.[/CONTENT]
# [CONTENT]Kristallar o'zining belgilangan tuzilishi va kristal qoidalari asosida xususiyatlarga ega bo'ladi. Ular o'rtacha shakli, ranglari, kuchini va to'g'ri joylashuvi bilan ajralib turadi.[/CONTENT]
# [SUBTITLE]Radioaktivlik: Atom yadrosining parchalanishi.[/SUBTITLE]
# [HEADING]Radioaktivlik: Atom yadrosining parchalanishi.[/HEADING]
# [CONTENT]Radioaktivlik - bu atom yadrosining spontan parchalanishi. Bu jarayon orqali atom yadro bo'lgan element energiya va subatomar zarrachalar yoki fotonlarga aylanadi.[/CONTENT]
# [CONTENT]Radioaktivlikning asosiy sababi atom yadrosining instabil bo'lishi va bu yadroni o'z-o'zidan zarrachalar va energiya yoki fotonlar bo'lishi talab qilinadi. Bu esa atom yadrosining parchalanishiga olib keladi.[/CONTENT]
# [CONTENT]Radioaktiv elementlar faqat bir necha yarim hayotiy zamon va yaqin zamin turi orqali parchalanadi. Bu jarayon energiya yoki zarracha emisiyasiga olib keladi va fotonlar yaratadi.[/CONTENT]
# [CONTENT]Radioaktivlik fazilatli davom etgay va bu elementlarning parchalanish tezligi yadroli elementiga qarab farqli bo'ladi. Ushbu jarayon radioisotoplarning ko'rsatkichi orqali baholanadi.[/CONTENT]
# [SUBTITLE]Atom spektrlari va elektron qobiqlarining energiya darajalari.[/SUBTITLE]
# [HEADING]Atom spektrlari va elektron qobiqlarining energiya darajalari.[/HEADING]
# [CONTENT]Atom spektrlari atomdan yoki molekuladan yoki ionni tark etish paytida emitrlgan yoki absorbatsiya qilingan noor sodir bo'lganda noorni daftarkoshini yoki noori yozadigan ikkita havolalardagi eroziydan iborat darvozadan tushadi.[/CONTENT]
# [CONTENT]Elektron qobiqlari atom muloqotini kuch-sutkasini ifodalovchi qonunlar orqali qayta qayta yaratiladi. Bu kuch va maskuniyat har bir elektron uchun taxminiy energiya qiymatlarini belgilaydi.[/CONTENT]
# [CONTENT]Atom energiya darajalari kvantlanish va elektronning muloqot parametrlari asosida belgilanadi. Bu energiya darajalari atomlarning optik va elektron spektrlarini tushunishga yordam beradi.[/CONTENT]
# [CONTENT]Elektron qobiqlarining spektrali atom spektrining individual tomonlari bo'lib, elektron qobiqlarining va energiya darajalarining o'zaro bog'liqlikni ko'rsatadi. Bu spektraldan foydalanib atomning xususiyatlarini tushunish mumkin.[/CONTENT]
# [SUBTITLE]Atomning fizikaviy va kimyoviy xususiyatlarini aniqlashda elektron konfiguratsiyasining ahamiyati.[/SUBTITLE]
# [HEADING]Atomning fizikaviy va kimyoviy xususiyatlarini aniqlashda elektron konfiguratsiyasining ahamiyati.[/HEADING]
# [CONTENT]Elektron konfiguratsiyasi atomdagi elektronlar sonini va ularning orbitalda joylashuvi haqida ma'lumot beradi.[/CONTENT]
# [CONTENT]Elektron konfiguratsiyasi atomning xususiyatlari, kimyoviy reaksiyalarni va kimyoviy aloqalarini aniqlashda juda muhim ahamiyatga ega.[/CONTENT]
# [CONTENT]Elektron konfiguratsiyasi atom uchun kimyoviy aloqalar, molekulalar va kimyoviy birliklarga o'z xususiyatlarini belgilab beradigan asosiy faktor hisoblanadi.[/CONTENT]
# [CONTENT]Atomning magnetik, elektrik va optik xususiyatlari elektron konfiguratsiyasiga bog'liq bo'lishi sababli, bu xususiyatlar atomlarni tanimasda o'ziga xos muhim ahamiyatga ega.[/CONTENT]
# [SUBTITLE]Atomlardagi o‘zaro ta’sir: Van der Vaals kuchlari va gidrogen bog‘lanishlari.[/SUBTITLE]
# [HEADING]Atomlardagi o‘zaro ta’sir: Van der Vaals kuchlari va gidrogen bog‘lanishlari.[/HEADING]
# [CONTENT]Van der Vaals kuchlari atomlarning molekulyalar o‘rtasida yuzaga keladigan jismlar bo‘lib, o‘zaro ta’sirning sababchisi sifatida faol bo‘ladi.[/CONTENT]
# [CONTENT]Gidrogen bog‘lanishlari bir atomning gidrogeni va boshqa atom yoki molekula bilan ta’sirlashish bo‘yicha yaratilgan kuchli kuvvatdir.[/CONTENT]
# [CONTENT]Atomlar o‘zaro ta’sir bo‘yicha Van der Vaals kuchlarida jismlar orasidagi moylik kuvvat ko‘rsatadigan va gidrogen bog‘lanishlari orqali elementlar o‘rtasida o‘zaro aloqani ta’minlaydi.[/CONTENT]
# [CONTENT]Van der Vaals kuchlari va gidrogen bog‘lanishlari atomlarning molekullyar orasidagi molekullar orasida yuzaga keladigan jismlar bo‘lib, atomlar o‘zaro ta’sir orqali bog‘lanadigan kuchli kuvvatlardir.[/CONTENT]
# [SUBTITLE]Insoniyat uchun atom tuzilishining ahamiyati: Energetika, tibbiyot va materialshunoslik.[/SUBTITLE]
# [HEADING]Insoniyat uchun atom tuzilishining ahamiyati: Energetika, tibbiyot va materialshunoslik.[/HEADING]
# [CONTENT]Atom tuzilishi energetikada yirik ahamiyatga ega: nukleer energiya, iste'mol qilish va uning qulayliklari insoniyat uchun zarur.[/CONTENT]
# [CONTENT]Atom tuzilishi tibbiyot sohasida muhim: radioaktiv terapiya, ronuvchi zaryadlar va boshqa xirurgik amaliyotlar tuzilgan tashqi holatlarda foydalaniladi.[/CONTENT]
# [CONTENT]Atom tuzilishi materialshunoslikda muhim rol o'ynaydi: nanotexnologiya, kompozit materiallar va boshqa yangiliklar atom tuzilishi asosida ishlab chiqiladi.[/CONTENT]"""
    
#     await generate_docx(sample_answer, "20240925")

# Run the async main function
# asyncio.run(main())

import sys
import asyncio

async def test():
    text = sys.stdin.read()
    print(text,"NImadirlar")
    timestampt = (text.split("{{")[0])
    txt = (text.split("{{")[1])
    output, error = await generate_docx(txt,timestampt)
    print(output, error,"Salom")


if __name__ == "__main__":
    asyncio.run(test())
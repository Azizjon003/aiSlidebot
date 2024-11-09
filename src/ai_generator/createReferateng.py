import io
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import asyncio

async def generate_docx(answer, time):
    doc = Document()

    async def split_tags(reply):
        pattern = r'\[(.*?)\](.*?)\[/\1\]'
        tags = re.findall(pattern, reply, re.DOTALL)
        return tags

    async def parse_response(tags_array):
        if not tags_array:
            raise IndexError
        
        # Add title and subtitle
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run("Ministry Of Higher Education, Science And Innovations Of The Republic Of Uzbekistan")
        run1.bold = True
        run1.font.size = Pt(24)

        # p2 = doc.add_paragraph()
        # p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # run2 = p2.add_run("Muhammad Al-Xorazmiy nomidagi Toshkent\nAxborot Texnologiyalari Universiteti")
        # run2.bold = True
        # run2.font.size = Pt(22)
        university_name = "Muhammad Al-Xorazmiy nomidagi Toshkent\nAxborot Texnologiyalari Universiteti"
        for item in tags_array:
            if item[0] == 'UNIVERSITY':
                university_name = item[1]
                break

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(university_name)
        run2.bold = True
        run2.font.size = Pt(22)

        # Add empty lines
        for _ in range(4):
            doc.add_paragraph()

        # Add "O'rnatilgan tizimlar fanidan"
        theme_text = "O'rnatilgan tizimlar fanidan"
        for item in tags_array:
            if item[0] == 'THEME':
                theme_text = item[1]
                break

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p.add_run(theme_text)
        run1.font.size = Pt(16)

        doc.add_paragraph()

        # Add "Mustaqil ish"
        p = doc.add_paragraph("Mustaqil ish")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(20)

        # Gather subtitles for the table of contents
        subtitless = []
        right_content = ""
        for item in tags_array:
            if item[0] == 'SUBTITLE':
                subtitless.append(item[1])
            if item[0] == 'RIGHTCONTENT':
                right_content = item[1]  # Store the right content

        subtitles = [subtitless[0],subtitless[1],subtitless[2],subtitless[subtitless.__len__()-1]]
 
        # Add "Toshkent 2024"
        if right_content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(right_content)
            run.font.size = Pt(14)

        for _ in range(4):
          doc.add_paragraph()
        
        

        doc.add_page_break()

        # Add "Reja" (Table of Contents)
        mundarija = doc.add_paragraph("Plans")
        mundarija.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mundarija.runs[0].bold = True
        mundarija.runs[0].font.size = Pt(14)

        doc.add_paragraph()

        # Generate the Table of Contents based on SUBTITLE
        data = []
        page_number = 3  # Start page numbering from 3 (or adjust as needed)
        for i, subtitle in enumerate(subtitles):
            data.append((f"{i + 1}.", subtitle, str(page_number)))
            page_number += 1  # Increment page number for each subtitle

        # Add Table of Contents
        for num, content, page in data:
            paragraph = doc.add_paragraph()
            paragraph.add_run(num + " ")
            paragraph.add_run(content)
            
            paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_ALIGNMENT.DECIMAL)
            paragraph.add_run().add_tab()
            
            paragraph.add_run(page)

        doc.add_page_break()

        # Add content based on the tags_array
        for item in tags_array:
            match (item[0]):
                case 'HEADING':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(item[1])
                    run.bold = True
                    run.font.size = Pt(14)
                case 'CONTENT':
                    p = doc.add_paragraph(item[1])
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                # case 'RIGHTCONTENT':
                #     # Add the RIGHTCONTENT section
                #     p = doc.add_paragraph()
                #     p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                #     run = p.add_run(item[1])
                #     run.font.size = Pt(12)
                case 'IMAGE':
                    # Image handling is commented out as it requires additional setup
                    pass

    async def find_title(tags_array):
        for item in tags_array:
            if item[0] == 'TITLE':
                return item[1]
        return "Untitled Document"

    reply_array = await split_tags(answer)
    await parse_response(reply_array)
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    docx_title = f"{await find_title(reply_array)}.docx"
    with open(f"output2{time}.docx", "wb") as f:
        f.write(docx_bytes)
    print(f"Document saved as output2{time}.docx")
    print(f"Done: {docx_title}")

    return docx_bytes, docx_title



import sys
import asyncio

async def test():
    text = sys.stdin.read()
    print(text,"NImadirlar")
    timestampt = (text.split("{{")[0])
    txt = (text.split("{{")[1])
    output, error = await generate_docx(txt,timestampt)
    print(output, error,"Salom")


if __name__ == "__main__":
    asyncio.run(test())